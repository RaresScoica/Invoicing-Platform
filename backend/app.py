import base64
import logging
import threading
import certifi
import os
import requests
import json
import time
import datetime
import smtplib

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from dotenv import load_dotenv
from email import encoders
from email.mime.base import MIMEBase
from email.mime.image import MIMEImage
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from datetime import datetime
from flask import Flask, redirect, render_template, request, jsonify, send_file, send_from_directory, session, url_for
from pymongo import MongoClient, UpdateOne
from pymongo.server_api import ServerApi

load_dotenv()

logging.basicConfig(level=logging.DEBUG)

frontend_folder = os.path.join(os.path.dirname(__file__), '../frontend')
app = Flask(__name__, template_folder=os.path.join(frontend_folder, 'templates'))

# if 'RENDER' in os.environ:
#     app.config['SECRET_KEY'] = process.getenv('SECRET_KEY')
#     uri = os.getenv('MONGO_URI')
#     MAIL_USER = os.getenv('MAIL_USER')
#     MAIL_PASSWORD = os.getenv('MAIL_PASSWORD')
# else:
app.config['SECRET_KEY'] = os.getenv('SECRET_KEY')
uri = os.getenv('MONGO_URI')
MAIL_USER = os.getenv('MAIL_USER')
MAIL_PASSWORD = os.getenv('MAIL_PASSWORD')

# Create a new client and connect to the server
client = MongoClient(uri, server_api=ServerApi('1'), tlsCAFile=certifi.where())

# Send a ping to confirm a successful connection
try:
    client.admin.command('ping')
    print("Pinged your deployment. You successfully connected to MongoDB!")
except Exception as e:
    print(e)

def watch_collection():
    db = client['EV_Stations']
    collection = db['transactions']
    # Watch the collection for new inserts
    with collection.watch() as stream:
        for change in stream:
            if change['operationType'] == 'insert':
                # Handle the new object here
                new_document = change['fullDocument']
                print("New document inserted: ", new_document)
                transactionId = new_document.get("TransactionID")
                email = "invoice@arsek.ro"
                generate_docx(transactionId, email)

# Start the Change Stream watcher in a separate thread
threading.Thread(target=watch_collection, daemon=True).start()

def changeConfig():
    db = client['EV_Stations']
    collection = db['transactions']

    # Filter for documents to update
    filter = { 'sn': 'C6E20CCC23CATETRVT' }

    try:
        print("Attempting to find documents...")
        documents = collection.find(filter)
        documents_list = list(documents)
        print(f"Documents found: {documents_list}")
        if not documents_list:
            print("No documents found matching the filter.")
        else:
            print(f"Found {len(documents_list)} document(s) matching the filter.")
    except Exception as e:
        print(f"Error finding documents: {e}")
        raise

    # Initialize the counter
    k = 1

    # Prepare bulk operations
    bulk_updates = []
    for doc in documents:
        update = { '$set': { 'nr': k } }
        bulk_updates.append(UpdateOne({ '_id': doc['_id'] }, update))
        k += 1

    # Execute the bulk operations
    if bulk_updates:
        result = collection.bulk_write(bulk_updates)
        print(f'Matched {result.matched_count} document(s) and modified {result.modified_count} document(s)')
    else:
        print('No documents found matching the filter.')

@app.route('/')
def index():
    session.clear()
    return render_template('index.html')

@app.route('/insuccess')
def insuccess():
    session.clear()
    return render_template('insuccess.html')

@app.route('/<int:transactionId>')
def transaction(transactionId):
    return render_template('index.html', transactionId=transactionId)

@app.route('/success')
def success():
    db = client['EV_Stations']
    collection = db['current_transaction']

    current = collection.find_one({"ID": "current"})
    email = current.get("email")
    transactionId = current.get("transactionId")

    generate_docx(transactionId, email)

    return render_template('success.html', email=email)

def generate_docx(transactionId, email):
    company_details = ""
    try:
        company_details = session.get('company_details')
    except:
        logging.debug("Nu e CUI")

    try:
        transactionId = int(transactionId)
    except Exception as e:
        return redirect(url_for('index'))

    db = client['EV_Stations']
    collectionTransactions = db['transactions']

    transactionDetails = collectionTransactions.find_one({"TransactionID": transactionId})
    print(transactionDetails)

    if transactionDetails['finalAmount'] != 0:
        if transactionDetails == None:
            return redirect(url_for('index'))
        
        nr = transactionDetails['nr']
        collectionStations = db['stations']
        station = collectionStations.find_one({"sn": transactionDetails["sn"]})
        series = station["series"]
        
        # Get the current date
        formatted_time = transactionDetails["StopTime"]
        # Parse the input string into a datetime object
        parsed_time = datetime.strptime(formatted_time, "%Y-%m-%dT%H:%M:%SZ")
        # Format the datetime object in the desired format
        current_date = parsed_time.strftime("%d/%m/%Y %H:%M")

        if 'RENDER' in os.environ:
            with open('/opt/render/project/src/frontend/images/dfg_logo.png', 'rb') as f:
                image_data = f.read()
        else:
            with open('../frontend/images/dfg_logo.png', 'rb') as f:
                image_data = f.read()

        # Convert image data to base64-encoded string
        base64_image = base64.b64encode(image_data).decode('utf-8')

        with app.app_context():
            # Render the HTML template for the invoice and pass session storage data
            html = render_template('invoice.html', series=series, nr=nr, image_data=base64_image, email=email, company_details=company_details, transactionDetails=transactionDetails, current_date=current_date)

        # Create the Word document
        doc = Document()
        doc.add_heading('Factura/Invoice', 0)

        # Add company details
        doc.add_heading('Furnizor/Seller', level=1)
        doc.add_paragraph('DFG ACTIVE IMOBILIARE SRL')
        doc.add_paragraph('CUI/Tax ID no: 15830118')
        doc.add_paragraph('Adresa/Adress: MUNICIPIUL BUCUREŞTI, SECTOR 5, STR. ION CREANGĂ, NR.7, CAMERA 3, ET.6, AP.25')
        doc.add_paragraph('Registrul comertului/Registration no: J40/13984/2003')

        if company_details:
            # Add customer details
            doc.add_heading('Cumparator/Customer', level=1)
            doc.add_paragraph(f'Denumire: {company_details["denumire"]}')
            doc.add_paragraph(f'CUI/Tax ID no: {company_details["cui"]}')
            doc.add_paragraph(f'Adresa/Adress: {company_details["adresa"]}')
            doc.add_paragraph(f'Registrul comertului/Registration no: {company_details["nrRegCom"]}')
            doc.add_paragraph(f'Email: {email}')
        else:
            # Add customer details
            doc.add_heading('Cumparator/Customer', level=1)
            doc.add_paragraph(f'Persoana fizica')

        # Add invoice details
        doc.add_heading('Detalii Factura/Invoice Details', level=2)
        doc.add_paragraph(f'Numar factura/Invoice no: {series}-{nr}')
        doc.add_paragraph(f'Data emiterii/Date of issue: {current_date}')
        doc.add_paragraph(f'Data livrarii/Date of delivery: {current_date}')

        # Add transaction table
        table = doc.add_table(rows=1, cols=7)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Denumire produs\nProduct name'
        hdr_cells[1].text = 'Pret per KWh\nPrice per KWh\n(RON/KWh)'
        hdr_cells[2].text = 'Cantitate\nQuantity\n(KWh)'
        hdr_cells[3].text = 'Pret fara TVA\nPrice without VAT\n(RON)'
        hdr_cells[4].text = 'Valoare TVA\nVAT Value(RON)'
        hdr_cells[5].text = 'TVA\nVAT\n(%)'
        hdr_cells[6].text = 'Pret cu TVA\nPrice with VAT\n(RON)'

        # Add a row for the transaction details
        unit_price = round(transactionDetails["kwPrice"] / 1.19, 2)
        pre_tax_amount = round(transactionDetails["preTaxAmount"] / 100, 2)
        quantity = round(pre_tax_amount / unit_price, 2)
        vat_value = round(pre_tax_amount * 0.19, 2)
        total_value = round(pre_tax_amount + vat_value, 2)
        row_cells = table.add_row().cells
        row_cells[0].text = 'Energie/Energy'
        row_cells[1].text = f'{unit_price}'
        row_cells[2].text = f'{quantity}'
        row_cells[3].text = f'{pre_tax_amount}'
        row_cells[4] .text = f'{vat_value}'
        row_cells[5].text = '19'
        row_cells[6].text = f'{total_value}'

        # Add footer
        section = doc.sections[0]
        footer = section.footer
        footer_paragraph = footer.paragraphs[0]
        footer_paragraph.text = "Email: office@solar.planners.ro\nTelefon/Phone Number: 0726323012\nWebsite: solar.planners.ro\n\nFactura creata de Solar Planners SRL/Invoice created by Solar Planners SRL"
        footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        if 'RENDER' in os.environ:
            # Save the Word document
            docx_filename = f"/opt/render/project/src/backend/facturi/factura_{series}_{nr}.docx"
        else:
            docx_filename = f"facturi/factura_{series}_{nr}.docx"
        
        doc.save(docx_filename)
        send_emails(docx_filename, transactionId, email, series, nr)

    else:
        print("This is a remote transaction so it doesn't require an invoice!")

def remove_alpha_chars(s):
    """Remove non-digit characters from a string."""
    return ''.join(filter(str.isdigit, s))

def fetch_anaf_data(cui, attempts=2):
    """Fetch company data from ANAF API with retry on empty 'found'."""
    url = "https://webservicesp.anaf.ro/PlatitorTvaRest/api/v8/ws/tva"
    headers = {'Content-Type': 'application/json'}
    data = [{"cui": int(remove_alpha_chars(cui)), "data": datetime.now().strftime("%Y-%m-%d")}]
    
    for attempt in range(attempts):
        response = requests.post(url, headers=headers, json=data)
        if response.status_code == 200:
            try:
                result = response.json()
                if result.get('found'):
                    return result
                else:
                    print("Empty 'found' array, retrying...")
                    time.sleep(1.5)  # Wait for 1.5 seconds before retrying
            except json.JSONDecodeError:
                print("Failed to decode JSON. Response was:")
                print(response.text)
                return {}
        else:
            print(f"Failed to fetch data. Status code: {response.status_code}, Response: {response.text}")
            return {}

    print("Maximum attempts reached, returning last response.")
    return {}

# Function to save JSON data to a temporary file
def save_json(data):
    if 'RENDER' in os.environ:
        file_location = "/opt/render/project/src/temp/anaf_response.json"
    else:
        file_location = "C:/Users/developer/Documents/ws-server/platform/temp/anaf_response.json"  # Define temporary file path
    with open(file_location, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=4)
    return file_location

@app.route('/send_email_and_cui', methods=['POST'])
def send_email_and_cui():
    if 'RENDER' in os.environ:
        file_path = "/opt/render/project/src/temp/anaf_response.json"
    else:
        file_path = '../temp/anaf_response.json'
    if os.path.exists(file_path):
        os.remove(file_path)

    data = request.json
    cui = data.get('cui')
    email = data.get('email')
    transactionId = data.get('transactionId')

    db = client['EV_Stations']
    collection = db['current_transaction']

    # Combine the $set updates into a single dictionary
    update = {'$set': {'email': email, 'transactionId': transactionId, 'cui': cui}}

    # Update the document
    collection.update_one({"ID": "current"}, update, upsert=True)
    
    if cui:
        cui_variable = cui
        anaf_response = fetch_anaf_data(cui_variable)

        if anaf_response:
            file_location = save_json(anaf_response)
            collectionTransactions = db['transactions']
            transactionId = int(transactionId)

            # Query to check if 'cui' field exists in the document
            query = {'TransactionID': transactionId, 'cui': {'$exists': False}}

            # Check if a document with the given TransactionID exists and 'cui' field is not present
            document = collectionTransactions.find_one(query)

            if document:
                if 'cui' in document:
                    # 'cui' field already exists, return index.html
                    print("CUI field already exists.")
                    return render_template('index.html')
                else:
                    # Perform the update if the 'cui' field does not exist
                    result = collectionTransactions.update_one({'TransactionID': transactionId}, {'$set': {'cui': cui}})
                    if result.modified_count > 0:
                        print(f"Document updated, 'cui' field set to {cui}.")
                    else:
                        print("Document found but no changes were made.")
            else:
                print("No document found with the given TransactionID.")
                return jsonify({'message': 'CUI not good'}), 400
            
            print(f"Data saved to {file_location}")
            return jsonify({'file_location': file_location})
        else:
            return render_template('index.html')

    return jsonify({'message': 'CUI not provided'}), 400

@app.route('/send_email', methods=['POST'])
def send_email():
    data = request.json
    email = data.get('email')
    transactionId = data.get('transactionId')

    db = client['EV_Stations']
    collection = db['current_transaction']

    # Combine the $set updates into a single dictionary
    update = {'$set': {'email': email, 'transactionId': transactionId}}

    # Update the document
    collection.update_one({"ID": "current"}, update, upsert=True)
    
    return jsonify({'message': 'Data Received'})

@app.route('/send_company_details', methods=['POST'])
def send_company_details():
    data = request.json
    email = data.get('email')
    company_details = data.get('company')

    # Store email and company details in the session
    session['email'] = email
    session['company_details'] = company_details

    return 'Company details received successfully'

@app.route('/get_temp_file/<path:filename>')
def get_temp_file(filename):
    if 'RENDER' in os.environ:
        file_path = "/opt/render/project/src/temp/anaf_response.json"
    else:
        file_path = '../temp/anaf_response.json'
    if os.path.exists(file_path):
        # Serve the temporary file to the client
        if 'RENDER' in os.environ:
            return send_from_directory('/opt/render/project/src/temp', filename)
        else:
            return send_from_directory('C:/Users/developer/Documents/ws-server/platform/temp', filename)
    else:
        print("CUI not found")
        return redirect(url_for('index'))
    
def send_emails(attachment_file, transactionId, email, series, nr):
    # Get SMTP credentials from environmental variables
    smtp_username = MAIL_USER
    smtp_password = MAIL_PASSWORD
    
    sender_email = "invoice@arsek.ro"

    #"mihai.sandu@electricplanners.ro" , "romulus@dfg.ro"
    receiver_emails = [email, "rares.goiceanu@arsek.ro", "invoice@arsek.ro"]
    receiver_emails_str = ', '.join(receiver_emails)
    subject = "Factura {}_{}".format(series, nr)
    
    # HTML content for the email body
    body = """
    <html>
        <head>
            <style>
                body {{
                    font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
                }}
                .container {{
                    background-color: #f2f2f2;
                    padding: 24px;
                    border-radius: 3px;
                    border: 2px solid #b0c4de;
                }}
                .logo {{
                    text-align: left;
                    margin-right: 10px;
                }}
                .logo a {{
                    text-decoration: none;
                }}
                .logo img {{
                    max-width: 20%;
                    pointer-events: none;
                    -webkit-user-select: none;
                    -moz-user-select: none;
                    -ms-user-select: none;
                    user-select: none;
                }}
                .content {{
                    background-color: #ffffff;
                    padding: 20px;
                    border-radius: 3px;
                    margin-top: 20px;
                    color: #333333;
                }}
                p {{
                    margin: 0;
                    line-height: 1.5;
                    font-size: 16px;
                }}
                .spacer1 {{
                    margin-bottom: 20px;
                }}
                .spacer2 {{
                    margin-bottom: 50px;
                }}
            </style>
        </head>
        <body>
            <div class="container">
                <div class="logo">
                    <a href="https://arsek.ro" target="_self">
                        
                    </a>
                </div>
                <div class="content">
                    <p>Buna ziua,</p>
                    <div class="spacer1"></div>
                    <p>Am atasat factura pentru tranzactia <b>{}</b>.</p>
                    <div class="spacer2"></div> 
                    <p>Multumim,</p>
                    <p>Echipa Arsek Software</p>
                </div>
            </div>
        </body>
    </html>
    """.format(transactionId)
    
    # image <img src="cid:logo" alt="Arsek Logo">

    # Create message container
    msg = MIMEMultipart()
    msg['From'] = sender_email
    msg['To'] = receiver_emails_str
    msg['Subject'] = subject

    # Attach body
    #msg.attach(MIMEBase('text', 'plain'))
    msg.attach(MIMEText(body, 'html'))

    if 'RENDER' in os.environ:
        # Attach the logo image
        with open("/opt/render/project/src/frontend/images/logo_nobg.png", 'rb') as f:
            logo = MIMEImage(f.read(), _subtype="svg+xml")
            logo.add_header('Content-ID', '<logo>')
            msg.attach(logo)
    else:
        # Attach the logo image
        with open("../frontend/images/logo_nobg.png", 'rb') as f:
            logo = MIMEImage(f.read(), _subtype="svg+xml")
            logo.add_header('Content-ID', '<logo>')
            msg.attach(logo)

    # Attach file
    print(attachment_file)
    attachment = open(attachment_file, "rb")
    p = MIMEBase('application', 'octet-stream')
    p.set_payload(attachment.read())
    encoders.encode_base64(p)
    # Get just the filename from the attachment_file path
    attachment_filename = os.path.basename(attachment_file)
    
    p.add_header('Content-Disposition', f'attachment; filename={attachment_filename}')
    msg.attach(p)

    # Connect to SMTP server and send email
    smtp_server = "smtp.hostinger.com"
    smtp_port = 587

    server = smtplib.SMTP(smtp_server, smtp_port)
    server.starttls()
    server.login(smtp_username, smtp_password)
    server.sendmail(sender_email, receiver_emails, msg.as_string())
    print("Email with {} succesfully sent to {}".format(transactionId, receiver_emails_str))
    server.quit()
    
if __name__ == '__main__':
    # changeConfig()
    port = int(os.getenv('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=True, use_reloader=False, threaded=False)
